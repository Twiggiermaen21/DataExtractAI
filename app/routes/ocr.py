import os
import json
import logging
import time
from flask import Blueprint, request, jsonify, current_app, send_from_directory
from werkzeug.utils import secure_filename

from app.services.ocr_pipeline import get_pipeline, unload_pipeline
from app.auth import require_auth

log = logging.getLogger(__name__)

ocr_bp = Blueprint('ocr', __name__)

IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp'}


def _request_id():
    return f"ocr-{int(time.time() * 1000)}"


def _file_size(path):
    try:
        return os.path.getsize(path)
    except OSError:
        return None


def _fields_summary(fields):
    if not isinstance(fields, dict):
        return {'type': type(fields).__name__}
    keys = list(fields.keys())
    non_empty = [key for key, value in fields.items() if value not in (None, '', [], {})]
    return {
        'keys_count': len(keys),
        'non_empty_count': len(non_empty),
        'keys': keys[:30],
    }


@ocr_bp.route('/api/extract_pdf_text', methods=['POST'])
@require_auth
def extract_pdf_text():
    """Wyciaga surowy tekst z PDF/DOCX bez wysylania do LLM."""
    rid = _request_id()
    log.info("[%s] extract_pdf_text start: content_length=%s files=%s", rid, request.content_length, list(request.files.keys()))

    if 'file' not in request.files:
        log.warning("[%s] extract_pdf_text rejected: missing file field", rid)
        return jsonify({'success': False, 'error': 'Brak pliku'}), 400

    file = request.files['file']
    if file.filename == '':
        log.warning("[%s] extract_pdf_text rejected: empty filename", rid)
        return jsonify({'success': False, 'error': 'Pusta nazwa pliku'}), 400

    filename = secure_filename(file.filename) or 'upload'
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)

    try:
        file.save(filepath)
        ext = os.path.splitext(filename)[1].lower()
        log.info("[%s] extract_pdf_text saved: filename=%s ext=%s size=%s path=%s", rid, filename, ext, _file_size(filepath), filepath)
        text = ''

        if ext == '.pdf':
            import fitz
            doc = fitz.open(filepath)
            text = "".join(page.get_text() for page in doc)
            page_count = doc.page_count
            doc.close()
            log.info("[%s] extract_pdf_text pdf read: pages=%s chars=%s", rid, page_count, len(text))
        elif ext in ('.docx', '.doc'):
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            log.info("[%s] extract_pdf_text docx read: paragraphs=%s chars=%s", rid, len(doc.paragraphs), len(text))
        else:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
            log.info("[%s] extract_pdf_text plain read: chars=%s", rid, len(text))

        text = text.strip()
        original_len = len(text)
        if original_len > 3000:
            text = text[:3000]

        log.info("[%s] extract_pdf_text done: original_chars=%s returned_chars=%s truncated=%s", rid, original_len, len(text), original_len > 3000)
        return jsonify({
            'success': True,
            'text': text,
            'filename': filename,
            'original_length': original_len,
            'truncated': original_len > 3000,
        })

    except Exception as e:
        log.exception("[%s] extract_pdf_text error: filename=%s", rid, filename)
        return jsonify({'success': False, 'error': str(e)}), 500


@ocr_bp.route('/api/process_ocr', methods=['POST'])
@require_auth
def process_ocr():
    """OCR - przetwarza pliki i zwraca wyekstrahowane dane."""
    rid = _request_id()
    started_at = time.monotonic()
    log.info(
        "[%s] process_ocr start: remote=%s content_length=%s form_keys=%s file_keys=%s",
        rid,
        request.headers.get('X-Forwarded-For', request.remote_addr),
        request.content_length,
        list(request.form.keys()),
        list(request.files.keys()),
    )

    if 'files' not in request.files:
        log.warning("[%s] process_ocr rejected: missing files field", rid)
        return jsonify({'success': False, 'error': 'Brak plikow'}), 400

    files = request.files.getlist('files')
    incoming_files = [
        {
            'filename': file.filename,
            'content_type': file.content_type,
            'content_length': getattr(file, 'content_length', None),
        }
        for file in files
    ]
    log.info("[%s] process_ocr incoming files: count=%s files=%s", rid, len(files), incoming_files)

    if not files or files[0].filename == '':
        log.warning("[%s] process_ocr rejected: empty files list or empty filename", rid)
        return jsonify({'success': False, 'error': 'Nie wybrano plikow'}), 400

    # Sprawdź czy frontend przesłał własne pola (fields) - mają priorytet nad szablonem
    custom_fields_raw = request.form.get('fields')
    custom_fields = None
    if custom_fields_raw:
        try:
            custom_fields = json.loads(custom_fields_raw)
            if not isinstance(custom_fields, list) or not all(isinstance(f, str) for f in custom_fields):
                log.warning("[%s] process_ocr invalid fields format, ignoring: type=%s", rid, type(custom_fields).__name__)
                custom_fields = None
            else:
                log.info("[%s] process_ocr custom fields from frontend: count=%s fields=%s", rid, len(custom_fields), custom_fields[:10])
        except (json.JSONDecodeError, TypeError) as e:
            log.warning("[%s] process_ocr fields JSON parse error: %s", rid, e)
            custom_fields = None

    template_name = request.form.get('template', 'wezwanie_do_zaplaty.html')
    template_path = os.path.join(current_app.root_path, '..', 'templates', 'documents', template_name)
    template_exists = os.path.exists(template_path)
    model_name = request.form.get('model')
    log.info(
        "[%s] process_ocr config: template=%s template_exists=%s custom_fields=%s model=%s upload_folder=%s output_folder=%s",
        rid,
        template_name,
        template_exists,
        len(custom_fields) if custom_fields else None,
        model_name or '<env/default>',
        current_app.config['UPLOAD_FOLDER'],
        current_app.config['OUTPUT_FOLDER'],
    )

    try:
        # Jeśli frontend przesłał custom fields - nie ładuj szablonu, użyj pól z frontu
        if custom_fields:
            pipeline = get_pipeline(template_path=None, model=model_name, custom_fields=custom_fields)
        else:
            pipeline = get_pipeline(template_path if template_exists else None, model=model_name)
        if pipeline is None:
            log.error("[%s] process_ocr pipeline unavailable", rid)
            return jsonify({'success': False, 'error': 'Nie mozna polaczyc z LM Studio'}), 500
        log.info(
            "[%s] process_ocr pipeline ready: class=%s model=%s api_url=%s fields_count=%s",
            rid,
            pipeline.__class__.__name__,
            getattr(pipeline, 'model', None),
            getattr(pipeline, 'api_url', None),
            len(getattr(pipeline, 'fields', []) or []),
        )
    except Exception as e:
        log.exception("[%s] process_ocr pipeline init error", rid)
        return jsonify({'success': False, 'error': f'Blad: {str(e)}'}), 500

    processed_files = []
    documents = []
    errors = []

    for index, file in enumerate(files, start=1):
        if file.filename == '':
            log.warning("[%s] process_ocr skip empty filename at index=%s", rid, index)
            continue

        original_filename = file.filename
        filename = secure_filename(original_filename) or f'upload_{index}'
        original_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
        file_started_at = time.monotonic()
        log.info(
            "[%s] file %s/%s start: original=%s stored=%s content_type=%s path=%s",
            rid,
            index,
            len(files),
            original_filename,
            filename,
            file.content_type,
            original_path,
        )

        try:
            save_started_at = time.monotonic()
            file.save(original_path)
            log.info(
                "[%s] file saved: filename=%s size=%s save_ms=%d",
                rid,
                filename,
                _file_size(original_path),
                int((time.monotonic() - save_started_at) * 1000),
            )

            predict_started_at = time.monotonic()
            ocr_output = pipeline.predict(original_path)
            log.info(
                "[%s] pipeline.predict done: filename=%s results=%s predict_ms=%d",
                rid,
                filename,
                len(ocr_output) if ocr_output is not None else None,
                int((time.monotonic() - predict_started_at) * 1000),
            )

            has_data = False
            for result_index, res in enumerate(ocr_output or [], start=1):
                extracted = getattr(res, 'extracted_data', None)
                log.info(
                    "[%s] OCR result: filename=%s result=%s text_chars=%s extracted=%s",
                    rid,
                    filename,
                    result_index,
                    len(getattr(res, 'text', '') or ''),
                    _fields_summary(extracted),
                )

                saved = res.save_to_json(save_path=current_app.config['OUTPUT_FOLDER'])
                if saved:
                    saved_name = os.path.basename(saved)
                    processed_files.append(saved_name)
                    log.info("[%s] OCR result saved: filename=%s json=%s size=%s", rid, filename, saved_name, _file_size(saved))
                else:
                    log.warning("[%s] OCR result save failed: filename=%s result=%s", rid, filename, result_index)

                if extracted:
                    documents.append({'filename': filename, 'fields': extracted})
                    has_data = True

            if not has_data:
                error = f'Model nie zwrocil danych dla pliku: {filename}'
                log.warning("[%s] file no extracted data: filename=%s", rid, filename)
                errors.append({'file': filename, 'error': error})

            log.info("[%s] file done: filename=%s has_data=%s elapsed_ms=%d", rid, filename, has_data, int((time.monotonic() - file_started_at) * 1000))

        except Exception as e:
            log.exception("[%s] OCR error for %s", rid, filename)
            errors.append({'file': filename, 'error': str(e)})

    unload_pipeline()
    log.info("[%s] process_ocr pipeline unloaded", rid)

    success = len(documents) > 0
    status_code = 200 if success else 422
    message = f'Przetworzono {len(processed_files)} plikow'
    error_details = '; '.join(
        f"{item.get('file', 'plik')}: {item.get('error', 'brak danych')}"
        for item in errors
    )
    error_message = None if success else (
        error_details or 'OCR nie zwrocil danych dla zadnego pliku.'
    )

    response_payload = {
        'success': success,
        'processed': processed_files,
        'documents': documents,
        'errors': errors,
        'message': message,
        'error': error_message,
    }
    log.info(
        "[%s] process_ocr response: status=%s success=%s processed=%s documents=%s errors=%s elapsed_ms=%d",
        rid,
        status_code,
        success,
        len(processed_files),
        len(documents),
        errors,
        int((time.monotonic() - started_at) * 1000),
    )
    log.debug("[%s] process_ocr response payload: %s", rid, json.dumps(response_payload, ensure_ascii=False, default=str))

    return jsonify(response_payload), status_code


@ocr_bp.route('/api/get_results')
@require_auth
def get_results():
    """Zwraca liste plikow JSON z folderu output."""
    output_folder = current_app.config['OUTPUT_FOLDER']
    if not os.path.exists(output_folder):
        return jsonify([])
    try:
        files = sorted([f for f in os.listdir(output_folder) if f.endswith('.json')], reverse=True)
        return jsonify(files)
    except Exception:
        return jsonify([])


@ocr_bp.route('/api/get_result/<filename>')
@require_auth
def get_result(filename):
    """Zwraca zawartosc pliku JSON."""
    try:
        path = os.path.join(current_app.config['OUTPUT_FOLDER'], filename)
        with open(path, 'r', encoding='utf-8') as f:
            return jsonify(json.load(f))
    except Exception as e:
        return jsonify({'error': str(e)}), 404


@ocr_bp.route('/input/<filename>')
@require_auth
def serve_input(filename):
    return send_from_directory(current_app.config['UPLOAD_FOLDER'], filename)


@ocr_bp.route('/saved/<filename>')
@require_auth
def serve_saved(filename):
    return send_from_directory(current_app.config['SAVED_FOLDER'], filename)
